# main.py
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import re
from docx import Document
import fitz  # PyMuPDF
import io
import os

# Flask 애플리케이션 객체를 'app'이라는 이름으로 생성합니다.
app = Flask(__name__)
CORS(app)  # 모든 도메인에서의 API 요청을 허용합니다.

# --- Health Check Endpoint ---
@app.route('/')
def health_check():
    """서버가 정상적으로 실행 중인지 확인하기 위한 경로입니다."""
    return "Backend server is running!"

# --- Helper & Core Logic Functions ---

def extract_korean(text):
    """주어진 텍스트에서 한글만 추출하여 공백으로 연결합니다."""
    korean_chars = re.findall(r'[가-힣]+', text)
    return " ".join(korean_chars)

def process_name_line_s1(line):
    """'*'로 시작하는 줄을 분석하여 이름 형식을 변경합니다."""
    if not line.startswith('*'):
        return None
    text = line.lstrip('*').strip()
    match = re.match(r'([가-힣\s]+)([a-zA-Z\s]+)', text)
    if not match:
        return None
    korean_part, english_part = match.group(1).strip(), match.group(2).strip()
    korean_names = korean_part.split()
    reformatted_korean = f"{korean_names[-1]}, {' '.join(korean_names[:-1])}" if len(korean_names) >= 2 else korean_part
    english_names = english_part.split()
    reformatted_english = f"{english_names[-1]}, {' '.join(english_names[:-1])}" if len(english_names) >= 2 else english_part
    return f"{reformatted_korean}{reformatted_english}"

def generate_regex_from_sample(sample):
    """
    사용자가 입력한 샘플 텍스트의 '구조'를 분석하여 범용적인 정규식을 생성합니다.
    로컬 버전의 모든 개선 사항이 적용되었습니다.
    """
    # 구분 기호를 찾기 위해 샘플 텍스트를 분석합니다.
    match = re.match(r'^\s*(.+?)\s*([^가-힣a-zA-Z0-9\s])(.*?)\2\s*$', sample)
    if not match:
        raise ValueError("입력 형식이 올바르지 않습니다. '용어<기호>원어<기호>' 형식으로 입력해주세요. 예: 라몬즈+Ramones+")
    
    delimiter = re.escape(match.group(2))
    
    # 첫 번째 부분(용어)에 대한 범용적인 정규식 패턴 (1~20자 제한)
    part1_pattern = r'[가-힣A-Za-z0-9\s.]{1,20}'

    # 두 번째 부분(원어)에 대한 범용적인 정규식 패턴 (쉼표로 구분된 다중 원어 지원)
    single_term_pattern = r'[A-Za-z0-9\s.\-()\u00C0-\u017F\u4E00-\u9FFF]+'
    part2_pattern = f'{single_term_pattern}(?:,\\s*{single_term_pattern})*'

    # 최종 정규식 생성
    return f"({part1_pattern}){delimiter}({part2_pattern}){delimiter}"

# --- Processing Logic for Each Tab ---

def handle_tab1(file, regex_pattern, decorator_str):
    """색인 추출 (Tab 1) 로직"""
    try:
        re.compile(regex_pattern)
    except re.error:
        raise ValueError("제공된 정규식 패턴이 유효하지 않습니다.")
    doc = Document(file)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    matches = re.findall(regex_pattern, full_text)
    if not matches:
        raise ValueError("정규식과 일치하는 내용을 찾지 못했습니다.")
    prefix, suffix = ('', '')
    if len(decorator_str) == 1:
        prefix = suffix = decorator_str
    elif len(decorator_str) >= 2:
        prefix, suffix = decorator_str[0], decorator_str[1]
    unique_formatted_matches = []
    seen = set()
    for match in matches:
        # re.findall은 캡처 그룹의 수에 따라 다른 결과를 반환할 수 있습니다.
        # 이 로직은 생성된 정규식(2개의 메인 캡처 그룹)에 맞춰져 있습니다.
        if isinstance(match, tuple) and len(match) >= 2:
            # part2_pattern의 non-capturing group 때문에 match[2]가 빈 값일 수 있습니다.
            # part1과 part2 전체를 사용합니다.
            part1, part2 = str(match[0]).strip(), str(match[1]).strip()
            if part1 and part2:
                formatted_item = f"{part1}{prefix}{part2}{suffix}"
            else:
                continue
        else:
            # 캡처 그룹이 하나이거나 없는 경우
            item_str = str(match).strip()
            if item_str:
                formatted_item = item_str
            else:
                continue
        
        if formatted_item and formatted_item not in seen:
            unique_formatted_matches.append(formatted_item)
            seen.add(formatted_item)
    if not unique_formatted_matches:
        raise ValueError("정규식과 일치하는 유효한 데이터를 찾지 못했습니다.")
    new_doc = Document()
    new_doc.add_heading("추출된 색인 목록", level=1)
    for item in unique_formatted_matches:
        new_doc.add_paragraph(item)
    
    file_stream = io.BytesIO()
    new_doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def handle_tab2(file):
    """인명 정리 (Tab 2) 로직"""
    source_doc = Document(file)
    result_doc = Document()
    table = result_doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text = '원본 내용', '변경 내용'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    for para in source_doc.paragraphs:
        original_text = para.text.strip()
        if not original_text: continue
        changed_text = process_name_line_s1(original_text)
        row_cells = table.add_row().cells
        row_cells[0].text, row_cells[1].text = original_text, changed_text if changed_text else ""
    file_stream = io.BytesIO()
    result_doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def handle_tab3(file):
    """한글 추출 (Tab 3) 로직"""
    source_doc = Document(file)
    if not source_doc.tables:
        raise ValueError("선택한 파일에 표(테이블)가 없습니다.")
    source_table = source_doc.tables[0]
    result_doc = Document()
    result_table = result_doc.add_table(rows=1, cols=3)
    result_table.style = 'Table Grid'
    hdr_cells = result_table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = '원본 내용 (*제거)', '한글 추출', '변경 내용'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    for row in source_table.rows[1:]:
        if len(row.cells) < 2: continue
        col1_text, col2_text = row.cells[0].text, row.cells[1].text
        row_cells = result_table.add_row().cells
        row_cells[0].text = col1_text.lstrip('*').strip()
        row_cells[1].text = extract_korean(col1_text)
        row_cells[2].text = col2_text
    file_stream = io.BytesIO()
    result_doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def handle_tab4(pdf_file, docx_file, page_range_str):
    """PDF 페이지 찾기 (Tab 4) 로직"""
    index_doc = Document(docx_file)
    if not index_doc.tables:
        raise ValueError("색인 파일에 테이블이 없습니다.")
    table = index_doc.tables[0]
    search_terms = [row.cells[1].text.strip() for i, row in enumerate(table.rows) if i > 0 and len(row.cells) >= 2 and row.cells[1].text.strip()]
    if not search_terms:
        raise ValueError("색인 테이블의 2열에서 검색할 단어를 찾을 수 없습니다.")
    pdf_doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    total_pages = pdf_doc.page_count
    page_range_iterator = range(total_pages)
    if page_range_str:
        match = re.match(r'^\s*(\d+)\s*-\s*(\d+)\s*$', page_range_str)
        if match:
            start_page, end_page = int(match.group(1)) - 1, int(match.group(2))
            if 0 <= start_page < end_page <= total_pages:
                page_range_iterator = range(start_page, end_page)
    normalized_terms_map = {term: "".join(term.split()) for term in search_terms}
    results = {term: [] for term in search_terms}
    for page_num in page_range_iterator:
        page = pdf_doc.load_page(page_num)
        page_text_normalized = "".join(page.get_text("text").split())
        if not page_text_normalized: continue
        for original_term, normalized_term in normalized_terms_map.items():
            if normalized_term in page_text_normalized and (page_num + 1) not in results[original_term]:
                results[original_term].append(page_num + 1)
    pdf_doc.close()
    new_doc = Document()
    new_doc.add_heading("PDF 색인 분석 결과", level=1)
    original_table = index_doc.tables[0]
    new_table = new_doc.add_table(rows=0, cols=len(original_table.columns) + 1)
    new_table.style = 'Table Grid'
    for i, row in enumerate(original_table.rows):
        new_row_cells = new_table.add_row().cells
        for j, cell in enumerate(row.cells):
            new_row_cells[j].text = cell.text
        if i == 0:
            new_row_cells[-1].text = "페이지"
        else:
            if len(row.cells) > 1:
                term_to_find = row.cells[1].text.strip()
                found_pages = results.get(term_to_find)
                new_row_cells[-1].text = ", ".join(map(str, sorted(found_pages))) if found_pages else "PDF에서 찾을 수 없음"
            else:
                new_row_cells[-1].text = ""
    file_stream = io.BytesIO()
    new_doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def handle_tab5(file):
    """테이블 조합 (Tab 5) 로직"""
    doc = Document(file)
    if not doc.tables:
        raise ValueError("업로드된 워드 파일에 테이블이 없습니다.")
    table = doc.tables[0]
    if len(table.columns) < 4:
        raise ValueError("테이블이 4열 이상으로 구성되어야 합니다.")
    processed_lines = []
    for row in table.rows[1:]:
        if len(row.cells) < 4: continue
        col2_text, col3_text, col4_text = row.cells[1].text.strip(), row.cells[2].text.strip(), row.cells[3].text.strip()
        # [수정됨] 로컬 버전과 동일하게 4칸 공백으로 변경
        if col2_text and not col3_text:
            processed_lines.append(f"{col2_text}    {col4_text}")
        elif col2_text and col3_text:
            processed_lines.append(f"{col3_text}    {col4_text}")
    if not processed_lines:
        raise ValueError("테이블에서 처리할 데이터를 찾지 못했습니다.")
    unique_sorted_lines = sorted(list(set(processed_lines)))
    new_doc = Document()
    new_doc.add_heading("테이블 정리 결과 (가나다순 정렬)", level=1)
    for line in unique_sorted_lines:
        new_doc.add_paragraph(line)
    file_stream = io.BytesIO()
    new_doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- API Endpoints ---

@app.route('/api/generate-regex', methods=['POST'])
def generate_regex_endpoint():
    """샘플 텍스트를 받아 정규식을 생성하는 API 엔드포인트입니다."""
    try:
        data = request.get_json()
        if not data or 'sample' not in data:
            return jsonify({"message": "샘플 텍스트가 필요합니다."}), 400
        
        sample_text = data['sample']
        generated_regex = generate_regex_from_sample(sample_text)
        
        return jsonify({"regex": generated_regex})

    except ValueError as e:
        return jsonify({"message": str(e)}), 400
    except Exception as e:
        print(f"An error occurred during regex generation: {e}")
        return jsonify({"message": "정규식 생성 중 서버 오류가 발생했습니다."}), 500

@app.route('/api/process/<task_name>', methods=['POST'])
def process_task(task_name):
    """파일 처리 요청을 받아 각 기능에 맞는 함수를 호출하는 메인 API 엔드포인트입니다."""
    try:
        if task_name == 'tab1':
            file = request.files.get('file')
            regex = request.form.get('regex')
            decorator = request.form.get('decorator')
            if not file: return jsonify({"message": "파일이 없습니다."}), 400
            if not regex: return jsonify({"message": "정규식이 없습니다."}), 400
            result_stream = handle_tab1(file, regex, decorator)
            filename = f"result_{task_name}.docx"

        elif task_name in ['tab2', 'tab3', 'tab5']:
            file = request.files.get('file')
            if not file: return jsonify({"message": "파일이 없습니다."}), 400
            
            if task_name == 'tab2': result_stream = handle_tab2(file)
            elif task_name == 'tab3': result_stream = handle_tab3(file)
            else: result_stream = handle_tab5(file) # tab5
            filename = f"result_{task_name}.docx"

        elif task_name == 'tab4':
            pdf_file = request.files.get('pdf_file')
            docx_file = request.files.get('docx_file')
            page_range = request.form.get('page_range')
            if not pdf_file or not docx_file: return jsonify({"message": "PDF 또는 DOCX 파일이 없습니다."}), 400
            result_stream = handle_tab4(pdf_file, docx_file, page_range)
            filename = f"result_{task_name}.docx"

        else:
            return jsonify({"message": "알 수 없는 작업입니다."}), 404

        return send_file(
            result_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except ValueError as e:
        return jsonify({"message": str(e)}), 400
    except Exception as e:
        print(f"An error occurred in task {task_name}: {e}")
        return jsonify({"message": "서버 내부 오류가 발생했습니다. 관리자에게 문의하세요."}), 500

# gunicorn 실행을 위한 설정
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=int(os.environ.get('PORT', 8080)))
